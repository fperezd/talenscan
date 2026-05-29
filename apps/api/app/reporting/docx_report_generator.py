"""Generador Word ejecutivo de Evaluación 360.

Replica visualmente la pantalla web mediante:
- Score card con shading de color por categoría
- Layouts de 3 columnas con celdas coloreadas (fortalezas/oportunidades/brechas)
- Badges como celdas inline
- Tabla de dimensiones con barras visuales
- Secciones con headers tipográficamente consistentes
"""

from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# --- Paleta (alineada con html_template.py) ---------------------------------
BRAND_BLUE = RGBColor(0x17, 0x7F, 0xC6)
BRAND_BLUE_DARK = RGBColor(0x0F, 0x5E, 0x94)
BRAND_BLACK = RGBColor(0x1D, 0x1D, 0x1B)
BRAND_GRAY = RGBColor(0x57, 0x57, 0x56)
BRAND_GRAY_LIGHT = RGBColor(0xC7, 0xC6, 0xC6)
EMERALD_700 = RGBColor(0x04, 0x78, 0x57)
ROSE_700 = RGBColor(0xBE, 0x12, 0x3C)
AMBER_700 = RGBColor(0xB4, 0x53, 0x09)
INDIGO_700 = RGBColor(0x43, 0x38, 0xCA)


# Hex shading sin '#' (formato Word)
SHADE_EMERALD = "ECFDF5"
SHADE_EMERALD_BORDER = "A7F3D0"
SHADE_INDIGO = "EEF2FF"
SHADE_INDIGO_BORDER = "C7D2FE"
SHADE_ROSE = "FFF1F2"
SHADE_ROSE_BORDER = "FECDD3"
SHADE_AMBER = "FFFBEB"
SHADE_AMBER_BORDER = "FDE68A"
SHADE_BLUE_SOFT = "E5F1FA"
SHADE_SLATE = "F8FAFC"
SHADE_SLATE_BORDER = "E2E8F0"
SHADE_WHITE = "FFFFFF"


def _category_color(category: str) -> RGBColor:
    lower = (category or "").lower()
    if "muy alto" in lower:
        return RGBColor(0x04, 0x78, 0x57)
    if "buen" in lower:
        return RGBColor(0x1D, 0x4E, 0xD8)
    if "parcial" in lower:
        return RGBColor(0xB4, 0x53, 0x09)
    if "bajo" in lower:
        return RGBColor(0x47, 0x55, 0x69)
    return RGBColor(0xBE, 0x12, 0x3C)


def _category_shade(category: str) -> str:
    lower = (category or "").lower()
    if "muy alto" in lower:
        return SHADE_EMERALD
    if "buen" in lower:
        return "EFF6FF"
    if "parcial" in lower:
        return SHADE_AMBER
    if "bajo" in lower:
        return "F1F5F9"
    return SHADE_ROSE


def _add_run(paragraph, text: str, *, bold=False, italic=False, size=10, color=BRAND_BLACK) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.font.color.rgb = color


def _set_cell_shading(cell, hex_color: str) -> None:
    """Aplica color de fondo a una celda."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_border(cell, *, color: str = "E2E8F0", size: int = 6) -> None:
    """Aplica borde fino a todos los lados de la celda."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:color"), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120) -> None:
    """Margen interno de la celda en twentieths of a point."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, value in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def _spacer(document: Document, points: int = 4) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(points)
    p.paragraph_format.space_before = Pt(0)


def _section_header(document: Document, text: str) -> None:
    eyebrow = document.add_paragraph()
    eyebrow.paragraph_format.space_before = Pt(10)
    eyebrow.paragraph_format.space_after = Pt(0)
    _add_run(eyebrow, "TALENSCAN", bold=True, size=7, color=BRAND_BLUE)
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    _add_run(title, text, bold=True, size=13, color=BRAND_BLACK)


def _build_hero(document: Document, context: dict[str, Any]) -> None:
    """Header con score card + info candidato + recomendación + badges."""
    score = int(context.get("score", 0) or 0)
    category = str(context.get("score_category", ""))
    cat_color = _category_color(category)
    cat_shade = _category_shade(category)

    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(3.4)
    table.columns[1].width = Cm(13.6)

    # --- Score card cell ---
    score_cell = table.cell(0, 0)
    score_cell.width = Cm(3.4)
    score_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_shading(score_cell, cat_shade)
    cat_color_hex = "{:02X}{:02X}{:02X}".format(cat_color[0], cat_color[1], cat_color[2])
    _set_cell_border(score_cell, color=cat_color_hex)
    _set_cell_margins(score_cell, top=180, bottom=180, left=120, right=120)
    score_cell.text = ""

    p_score = score_cell.paragraphs[0]
    p_score.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_score.paragraph_format.space_after = Pt(0)
    run = p_score.add_run(str(score))
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = cat_color
    run.font.name = "Calibri"

    p_of = score_cell.add_paragraph()
    p_of.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_of.paragraph_format.space_after = Pt(0)
    _add_run(p_of, "/ 100", size=7, color=BRAND_GRAY)

    # --- Info cell ---
    info_cell = table.cell(0, 1)
    info_cell.width = Cm(13.6)
    info_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _set_cell_shading(info_cell, SHADE_WHITE)
    _set_cell_border(info_cell, color="E2E8F0")
    _set_cell_margins(info_cell, top=140, bottom=140, left=180, right=180)
    info_cell.text = ""

    eyebrow_p = info_cell.paragraphs[0]
    eyebrow_p.paragraph_format.space_after = Pt(2)
    _add_run(eyebrow_p, "EVALUACIÓN 360 TALENSCAN", bold=True, size=7, color=BRAND_BLUE)

    name_p = info_cell.add_paragraph()
    name_p.paragraph_format.space_after = Pt(2)
    _add_run(name_p, str(context.get("candidate_name", "")), bold=True, size=18, color=BRAND_BLACK)

    role_p = info_cell.add_paragraph()
    role_p.paragraph_format.space_after = Pt(4)
    _add_run(
        role_p,
        f"{context.get('target_role', '—')}  ·  {context.get('client_name', '—')}",
        size=10,
        color=BRAND_GRAY,
    )

    rec_p = info_cell.add_paragraph()
    rec_p.paragraph_format.space_after = Pt(6)
    _add_run(rec_p, str(context.get("recommendation", "")), size=10, color=BRAND_BLACK)

    # Badges row
    badges_p = info_cell.add_paragraph()
    badges_p.paragraph_format.space_after = Pt(0)
    _add_badge(badges_p, category, cat_color, cat_shade)
    critical = list(context.get("critical_gaps_detailed") or [])
    if critical:
        _add_badge(badges_p, f"  {len(critical)} brecha{'s' if len(critical) > 1 else ''} crítica{'s' if len(critical) > 1 else ''}  ", ROSE_700, SHADE_ROSE)
    else:
        _add_badge(badges_p, "  Sin brechas críticas  ", EMERALD_700, SHADE_EMERALD)
    opportunities = list(context.get("opportunities") or [])
    if opportunities:
        _add_badge(badges_p, f"  {len(opportunities)} oportunidad{'es' if len(opportunities) > 1 else ''} transferible{'s' if len(opportunities) > 1 else ''}  ", INDIGO_700, SHADE_INDIGO)


def _add_badge(paragraph, text: str, fg: RGBColor, shade_hex: str) -> None:
    """Badge inline: usamos highlight color simulado con shading no es posible inline,
    así que aplicamos color de fondo en el run via XML."""
    run = paragraph.add_run(f"  {text}  ")
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = fg
    run.font.name = "Calibri"
    # Background shading on the run
    r_pr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), shade_hex)
    r_pr.append(shd)
    spacer = paragraph.add_run("  ")
    spacer.font.size = Pt(8)


def _info_card(document: Document, body_paragraphs: list[tuple[str, Any]],
               shade_hex: str = SHADE_BLUE_SOFT, border_hex: str = "177FC6") -> None:
    """Renderiza un bloque destacado (1 celda)."""
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _set_cell_shading(cell, shade_hex)
    _set_cell_border(cell, color=border_hex, size=6)
    _set_cell_margins(cell, top=160, bottom=160, left=200, right=200)
    cell.text = ""
    first = True
    for kind, content in body_paragraphs:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(4)
        if kind == "eyebrow":
            _add_run(p, content, bold=True, size=7, color=BRAND_BLUE)
        elif kind == "title":
            _add_run(p, content, bold=True, size=14, color=BRAND_BLACK)
        elif kind == "quote":
            _add_run(p, content, bold=True, size=11, color=BRAND_BLACK)
        elif kind == "label":
            _add_run(p, content, bold=True, size=7, color=BRAND_GRAY)
        elif kind == "text":
            _add_run(p, content, size=10, color=BRAND_BLACK)
        elif kind == "muted":
            _add_run(p, content, size=10, color=BRAND_GRAY)


def _meta_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=len(rows), cols=2)
    table.autofit = False
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(12)
    for index, (label, value) in enumerate(rows):
        c_label = table.cell(index, 0)
        c_value = table.cell(index, 1)
        c_label.width = Cm(5)
        c_value.width = Cm(12)
        _set_cell_margins(c_label, top=40, bottom=40, left=80, right=80)
        _set_cell_margins(c_value, top=40, bottom=40, left=80, right=80)
        c_label.text = ""
        c_value.text = ""
        _add_run(c_label.paragraphs[0], label.upper(), bold=True, size=8, color=BRAND_GRAY)
        _add_run(c_value.paragraphs[0], value or "—", size=10, color=BRAND_BLACK)
        c_label.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c_value.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _three_col_layout(
    document: Document,
    columns: list[dict[str, Any]],
) -> None:
    """Genera un layout de hasta 3 columnas con celdas coloreadas.

    columns: lista de dicts con keys:
      - shade: hex de fondo
      - border: hex de borde
      - header_text, header_color
      - intro (opcional)
      - items: lista de dicts {title, detail, evidence, impact, mitigation}
      - chips_label, chips (opcional)
    """
    if not columns:
        return
    table = document.add_table(rows=1, cols=len(columns))
    table.autofit = False
    total_cm = 17.0
    col_width = Cm(total_cm / len(columns))
    for c_idx, col_data in enumerate(columns):
        cell = table.cell(0, c_idx)
        cell.width = col_width
        _set_cell_shading(cell, col_data["shade"])
        _set_cell_border(cell, color=col_data["border"], size=6)
        _set_cell_margins(cell, top=140, bottom=140, left=140, right=140)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        cell.text = ""

        # Header
        hdr_p = cell.paragraphs[0]
        hdr_p.paragraph_format.space_after = Pt(2)
        _add_run(hdr_p, col_data["header_text"], bold=True, size=10, color=col_data["header_color"])

        if col_data.get("intro"):
            intro_p = cell.add_paragraph()
            intro_p.paragraph_format.space_after = Pt(4)
            _add_run(intro_p, col_data["intro"], italic=True, size=8, color=col_data["header_color"])

        for item in col_data.get("items", []):
            # Sub-card (white)
            sep = cell.add_paragraph()
            sep.paragraph_format.space_after = Pt(2)
            title_p = cell.add_paragraph()
            title_p.paragraph_format.space_after = Pt(1)
            _add_run(title_p, item.get("title", ""), bold=True, size=9, color=BRAND_BLACK)
            if item.get("detail"):
                d_p = cell.add_paragraph()
                d_p.paragraph_format.space_after = Pt(1)
                _add_run(d_p, item["detail"], size=8, color=BRAND_GRAY)
            if item.get("impact"):
                i_p = cell.add_paragraph()
                i_p.paragraph_format.space_after = Pt(1)
                _add_run(i_p, "Impacto: ", bold=True, size=8, color=ROSE_700)
                _add_run(i_p, item["impact"], size=8, color=ROSE_700)
            if item.get("mitigation"):
                m_p = cell.add_paragraph()
                m_p.paragraph_format.space_after = Pt(1)
                _add_run(m_p, "Mitigación: ", bold=True, size=8, color=BRAND_BLACK)
                _add_run(m_p, item["mitigation"], size=8, color=BRAND_GRAY)
            if item.get("evidence"):
                e_p = cell.add_paragraph()
                e_p.paragraph_format.space_after = Pt(1)
                _add_run(e_p, "Evidencia: ", bold=True, italic=True, size=7, color=col_data["header_color"])
                _add_run(e_p, item["evidence"], italic=True, size=7, color=col_data["header_color"])

        if col_data.get("chips"):
            sep = cell.add_paragraph()
            sep.paragraph_format.space_after = Pt(2)
            _add_run(sep, col_data.get("chips_label", "Habilidades").upper(), bold=True, size=7, color=col_data["header_color"])
            chip_p = cell.add_paragraph()
            chip_p.paragraph_format.space_after = Pt(0)
            for chip_text in col_data["chips"]:
                chip_run = chip_p.add_run(f" {chip_text} ")
                chip_run.font.size = Pt(7)
                chip_run.font.color.rgb = col_data["header_color"]
                chip_run.font.name = "Calibri"
                r_pr = chip_run._element.get_or_add_rPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "FFFFFF")
                r_pr.append(shd)
                chip_p.add_run("  ")


def _trajectory_grid(document: Document, trajectory: dict[str, Any]) -> None:
    cells_data = []
    for label, key in (("Estabilidad", "tenure_stability"), ("Fase actual", "current_phase"), ("Progresión", "progression")):
        value = trajectory.get(key)
        if value:
            cells_data.append((label, str(value)))
    if not cells_data:
        return
    table = document.add_table(rows=1, cols=len(cells_data))
    table.autofit = False
    for idx, (label, value) in enumerate(cells_data):
        cell = table.cell(0, idx)
        cell.width = Cm(17 / len(cells_data))
        _set_cell_shading(cell, SHADE_SLATE)
        _set_cell_border(cell, color=SHADE_SLATE_BORDER)
        _set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
        cell.text = ""
        label_p = cell.paragraphs[0]
        _add_run(label_p, label.upper(), bold=True, size=7, color=BRAND_GRAY)
        value_p = cell.add_paragraph()
        _add_run(value_p, value, bold=True, size=10, color=BRAND_BLACK)


def _dimension_table(document: Document, dimensions: list[dict[str, Any]]) -> None:
    if not dimensions:
        return
    table = document.add_table(rows=1 + len(dimensions), cols=4)
    table.autofit = False
    widths = [Cm(6.5), Cm(2.8), Cm(2.6), Cm(5.1)]
    for col_idx, w in enumerate(widths):
        table.columns[col_idx].width = w

    # Header
    headers = ["DIMENSIÓN", "SCORE", "EVIDENCIA", "JUSTIFICACIÓN"]
    for col_idx, label in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.width = widths[col_idx]
        _set_cell_shading(cell, SHADE_SLATE)
        _set_cell_border(cell, color=SHADE_SLATE_BORDER)
        _set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
        cell.text = ""
        _add_run(cell.paragraphs[0], label, bold=True, size=7, color=BRAND_GRAY)

    # Rows
    for r_idx, dim in enumerate(dimensions, start=1):
        score = dim.get("score") or 0
        max_score = dim.get("max_score") or 0
        evidence = dim.get("evidence_level") or "—"
        rationale = dim.get("rationale") or "—"
        name = dim.get("dimension") or "Dimensión"

        cells = [table.cell(r_idx, ci) for ci in range(4)]
        for ci, c in enumerate(cells):
            c.width = widths[ci]
            _set_cell_shading(c, SHADE_WHITE)
            _set_cell_border(c, color="F1F5F9")
            _set_cell_margins(c, top=80, bottom=80, left=100, right=100)
            c.text = ""

        _add_run(cells[0].paragraphs[0], str(name), bold=True, size=9, color=BRAND_BLACK)
        score_p = cells[1].paragraphs[0]
        score_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(score_p, f"{score} / {max_score}", bold=True, size=9, color=BRAND_BLACK)

        ev_lower = str(evidence).lower()
        ev_color = EMERALD_700 if "alta" in ev_lower else AMBER_700 if "media" in ev_lower else BRAND_GRAY
        ev_shade = SHADE_EMERALD if "alta" in ev_lower else SHADE_AMBER if "media" in ev_lower else SHADE_SLATE
        ev_p = cells[2].paragraphs[0]
        ev_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_badge(ev_p, str(evidence), ev_color, ev_shade)

        _add_run(cells[3].paragraphs[0], str(rationale), size=8, color=BRAND_GRAY)


def _bullet_list(document: Document, items: list[str], color: RGBColor = BRAND_BLACK) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        _add_run(p, str(item), size=10, color=color)


def _signals_block(document: Document, signals: list[dict[str, Any]]) -> None:
    for signal in signals:
        indicator = str(signal.get("indicator", "Neutral")).lower()
        if "positive" in indicator:
            shade = SHADE_EMERALD
            border = SHADE_EMERALD_BORDER
            color = EMERALD_700
        elif "risk" in indicator:
            shade = SHADE_ROSE
            border = SHADE_ROSE_BORDER
            color = ROSE_700
        else:
            shade = SHADE_SLATE
            border = SHADE_SLATE_BORDER
            color = BRAND_GRAY

        table = document.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Cm(13)
        table.columns[1].width = Cm(4)
        signal_cell = table.cell(0, 0)
        ind_cell = table.cell(0, 1)
        for c in (signal_cell, ind_cell):
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_shading(c, shade)
            _set_cell_border(c, color=border)
            _set_cell_margins(c, top=80, bottom=80, left=140, right=140)
            c.text = ""
        _add_run(signal_cell.paragraphs[0], str(signal.get("signal", "")), size=10, color=BRAND_BLACK)
        ind_p = ind_cell.paragraphs[0]
        ind_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_run(ind_p, str(signal.get("indicator", "Neutral")).upper(), bold=True, size=7, color=color)


def _interview_questions_block(document: Document, questions: list[dict[str, Any]]) -> None:
    for index, q in enumerate(questions, start=1):
        table = document.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Cm(1.2)
        table.columns[1].width = Cm(15.8)

        num_cell = table.cell(0, 0)
        body_cell = table.cell(0, 1)
        num_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        body_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _set_cell_shading(num_cell, SHADE_BLUE_SOFT)
        _set_cell_border(num_cell, color="C7E0EF")
        _set_cell_shading(body_cell, SHADE_SLATE)
        _set_cell_border(body_cell, color=SHADE_SLATE_BORDER)
        _set_cell_margins(num_cell, top=140, bottom=140, left=80, right=80)
        _set_cell_margins(body_cell, top=120, bottom=120, left=160, right=160)
        num_cell.text = ""
        body_cell.text = ""

        n_p = num_cell.paragraphs[0]
        n_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(n_p, str(index), bold=True, size=11, color=BRAND_BLUE)

        q_p = body_cell.paragraphs[0]
        _add_run(q_p, str(q.get("question", "")), size=10, color=BRAND_BLACK)

        meta_parts = []
        if q.get("objective"):
            meta_parts.append(("Objetivo: ", str(q["objective"])))
        priority = str(q.get("priority") or "").strip()

        if meta_parts or priority:
            meta_p = body_cell.add_paragraph()
            meta_p.paragraph_format.space_before = Pt(2)
            for label, val in meta_parts:
                _add_run(meta_p, label, bold=True, size=8, color=BRAND_BLACK)
                _add_run(meta_p, val, size=8, color=BRAND_GRAY)
            if priority:
                pri_lower = priority.lower()
                if "alta" in pri_lower:
                    pri_shade, pri_color = SHADE_ROSE, ROSE_700
                elif "media" in pri_lower:
                    pri_shade, pri_color = SHADE_AMBER, AMBER_700
                else:
                    pri_shade, pri_color = SHADE_SLATE, BRAND_GRAY
                _add_badge(meta_p, f"  Prioridad {priority}  ", pri_color, pri_shade)


def build_docx_report(context: dict[str, Any]) -> bytes:
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    _build_hero(document, context)
    _spacer(document, 8)

    # --- Tesis de talento ---
    talent_thesis = context.get("talent_thesis")
    if talent_thesis:
        body = [
            ("eyebrow", "TESIS DE TALENTO"),
            ("quote", str(talent_thesis)),
        ]
        if context.get("differentiation"):
            body.append(("label", "DIFERENCIADOR"))
            body.append(("text", str(context["differentiation"])))
        _info_card(document, body, shade_hex=SHADE_BLUE_SOFT, border_hex="177FC6")
        _spacer(document, 6)

    # --- Resumen ejecutivo ---
    _section_header(document, "Resumen ejecutivo")
    summary_p = document.add_paragraph()
    _add_run(summary_p, str(context.get("summary", "Sin resumen disponible.")), size=10, color=BRAND_BLACK)

    if context.get("final_verdict"):
        verdict_card_body = [
            ("label", "VEREDICTO FINAL"),
            ("text", str(context["final_verdict"])),
        ]
        _info_card(document, verdict_card_body, shade_hex=SHADE_SLATE, border_hex=SHADE_SLATE_BORDER)

    # --- Mandato meta ---
    _section_header(document, "Información del mandato")
    meta_rows = [
        ("Cliente", str(context.get("client_name", "—"))),
        ("Mandato", str(context.get("mandate_title", "—"))),
        ("Cargo objetivo", str(context.get("target_role", "—"))),
        ("Perfil objetivo", str(context.get("position_spec_title", "—"))),
    ]
    if context.get("candidate_position"):
        meta_rows.append(("Cargo actual", str(context["candidate_position"])))
    if context.get("candidate_company"):
        meta_rows.append(("Empresa actual", str(context["candidate_company"])))
    if context.get("candidate_email"):
        meta_rows.append(("Email", str(context["candidate_email"])))
    if context.get("candidate_phone"):
        meta_rows.append(("Teléfono", str(context["candidate_phone"])))
    _meta_table(document, meta_rows)

    # --- 3 columnas: Fortalezas / Oportunidades / Brechas ---
    strengths_detailed = list(context.get("strengths_detailed") or [])
    opportunities = list(context.get("opportunities") or [])
    critical_gaps_detailed = list(context.get("critical_gaps_detailed") or [])
    transferable_skills = list(context.get("transferable_skills") or [])

    if strengths_detailed or opportunities or critical_gaps_detailed:
        _section_header(document, "Calce 360 — Fortalezas, oportunidades y brechas")
        columns = []
        if strengths_detailed:
            columns.append({
                "shade": SHADE_EMERALD,
                "border": SHADE_EMERALD_BORDER,
                "header_text": f"Fortalezas calzadas · {len(strengths_detailed)}",
                "header_color": EMERALD_700,
                "items": [
                    {"title": str(s.get("title", "Fortaleza")), "detail": s.get("detail"), "evidence": s.get("evidence")}
                    for s in strengths_detailed if isinstance(s, dict)
                ],
            })
        if opportunities:
            opp_items = [
                {"title": str(o.get("title", "Oportunidad")), "detail": o.get("detail"), "evidence": o.get("evidence")}
                for o in opportunities if isinstance(o, dict)
            ]
            columns.append({
                "shade": SHADE_INDIGO,
                "border": SHADE_INDIGO_BORDER,
                "header_text": f"Oportunidades transferibles · {len(opportunities)}",
                "header_color": INDIGO_700,
                "intro": "Fortalezas que el cliente puede no estar valorando.",
                "items": opp_items,
                "chips_label": "Habilidades transferibles",
                "chips": transferable_skills,
            })
        if critical_gaps_detailed:
            columns.append({
                "shade": SHADE_ROSE,
                "border": SHADE_ROSE_BORDER,
                "header_text": f"Brechas críticas · {len(critical_gaps_detailed)}",
                "header_color": ROSE_700,
                "items": [
                    {
                        "title": str(g.get("requirement", "Brecha")),
                        "impact": g.get("impact"),
                        "mitigation": g.get("mitigation"),
                        "evidence": g.get("evidence"),
                    }
                    for g in critical_gaps_detailed if isinstance(g, dict)
                ],
            })
        _three_col_layout(document, columns)

    # --- Trayectoria ---
    trajectory = context.get("career_trajectory") or {}
    if isinstance(trajectory, dict) and any(trajectory.values()):
        _section_header(document, "Trayectoria de carrera")
        _trajectory_grid(document, trajectory)
        if trajectory.get("narrative"):
            narr_p = document.add_paragraph()
            narr_p.paragraph_format.space_before = Pt(4)
            _add_run(narr_p, str(trajectory["narrative"]), size=10, color=BRAND_GRAY)

    # --- Score por dimensión ---
    dimensions = list(context.get("dimensions") or [])
    if dimensions:
        _section_header(document, "Score por dimensión")
        _dimension_table(document, dimensions)

    # --- Cultural fit ---
    cultural = list(context.get("cultural_fit_signals") or [])
    if cultural:
        _section_header(document, "Señales de cultural fit")
        _signals_block(document, cultural)

    # --- Debilidades ---
    weaknesses = [str(x) for x in (context.get("weaknesses") or [])]
    if weaknesses:
        _section_header(document, "Debilidades manejables")
        _bullet_list(document, weaknesses)

    # --- Riesgos ---
    risks_detailed = list(context.get("risks_detailed") or [])
    if risks_detailed:
        _section_header(document, "Riesgos a validar")
        for risk in risks_detailed:
            if not isinstance(risk, dict):
                continue
            risk_table = document.add_table(rows=1, cols=1)
            cell = risk_table.cell(0, 0)
            _set_cell_shading(cell, SHADE_AMBER)
            _set_cell_border(cell, color=SHADE_AMBER_BORDER)
            _set_cell_margins(cell, top=100, bottom=100, left=160, right=160)
            cell.text = ""
            r_p = cell.paragraphs[0]
            _add_run(r_p, str(risk.get("risk", "")), size=10, color=BRAND_BLACK)
            if risk.get("validation"):
                v_p = cell.add_paragraph()
                _add_run(v_p, "Cómo validar: ", bold=True, size=8, color=AMBER_700)
                _add_run(v_p, str(risk["validation"]), size=8, color=AMBER_700)

    # --- Red flags ---
    red_flags = [str(x) for x in (context.get("red_flags") or [])]
    if red_flags:
        _section_header(document, "Red flags")
        _bullet_list(document, red_flags, color=ROSE_700)

    # --- Preguntas para entrevista ---
    questions_detailed = list(context.get("interview_questions_detailed") or [])
    if questions_detailed:
        _section_header(document, "Preguntas sugeridas para entrevista")
        _interview_questions_block(document, questions_detailed)

    # --- Reference check ---
    references = [str(x) for x in (context.get("reference_check_focus") or [])]
    if references:
        _section_header(document, "Foco en validación de referencias")
        _bullet_list(document, references)

    # --- Onboarding ---
    onboarding = [str(x) for x in (context.get("onboarding_considerations") or [])]
    if onboarding:
        _section_header(document, "Consideraciones de onboarding")
        _bullet_list(document, onboarding)

    # --- Compensación ---
    if context.get("compensation_signals"):
        _section_header(document, "Señales de compensación")
        comp_p = document.add_paragraph()
        _add_run(comp_p, str(context["compensation_signals"]), size=10, color=BRAND_GRAY)

    # --- Evidencia citada ---
    evidence = [str(x) for x in (context.get("evidence") or [])]
    if evidence:
        _section_header(document, "Evidencia citada del perfil")
        for item in evidence:
            table_ev = document.add_table(rows=1, cols=1)
            cell = table_ev.cell(0, 0)
            _set_cell_shading(cell, SHADE_SLATE)
            _set_cell_border(cell, color="F1F5F9")
            _set_cell_margins(cell, top=80, bottom=80, left=160, right=160)
            cell.text = ""
            _add_run(cell.paragraphs[0], f"“{item}”", italic=True, size=9, color=BRAND_GRAY)

    # --- Traceability footer ---
    _spacer(document, 8)
    foot_table = document.add_table(rows=1, cols=1)
    foot_cell = foot_table.cell(0, 0)
    _set_cell_shading(foot_cell, SHADE_SLATE)
    _set_cell_border(foot_cell, color=SHADE_SLATE_BORDER)
    _set_cell_margins(foot_cell, top=120, bottom=120, left=160, right=160)
    foot_cell.text = ""
    f_p = foot_cell.paragraphs[0]
    _add_run(f_p, "NOTA DE TRAZABILIDAD\n", bold=True, size=7, color=BRAND_GRAY)
    _add_run(f_p, str(context.get("traceability_note", "")), size=7, color=BRAND_GRAY)
    f2 = foot_cell.add_paragraph()
    _add_run(
        f2,
        f"Modelo: {context.get('model_version', '')} · Prompt: {context.get('prompt_version', '')}",
        size=7,
        color=BRAND_GRAY,
    )

    output = BytesIO()
    document.save(output)
    return output.getvalue()
