from io import BytesIO

from docx import Document


def _create_docx_bytes() -> bytes:
    document = Document()
    document.add_paragraph("Juan Perez")
    document.add_paragraph("Gerente Comercial - Empresa Demo")
    document.add_paragraph("12 anos de experiencia en tecnologia y SaaS.")
    document.add_paragraph("Logro: crecimiento de ingresos +35% y mejora de conversion.")
    document.add_paragraph("Herramientas: Salesforce, Power BI, Excel, SQL.")
    document.add_paragraph("Educacion: Ingenieria Comercial, MBA.")
    document.add_paragraph("Idioma: Ingles avanzado.")
    content = BytesIO()
    document.save(content)
    return content.getvalue()


def _create_mandate(client) -> int:
    payload = {
        "client_name": "Cliente PR4",
        "search_title": "Busqueda Head Comercial",
        "target_role": "Head Comercial",
        "must_have_requirements": ["Liderazgo comercial", "Gestion de pipeline"],
        "nice_to_have_requirements": ["MBA"],
        "main_responsibilities": ["Dirigir estrategia de ventas", "Liderar equipo"],
        "expected_results": ["Aumento de ingresos", "Mejora de conversion"],
        "target_industries": ["Tecnologia", "SaaS"],
        "status": "Activo",
    }
    response = client.post("/api/mandatos", json=payload)
    assert response.status_code == 201
    return response.json()["id"]


def test_candidates_documents_profiles_and_evaluations_flow(client) -> None:
    candidate_response = client.post(
        "/api/candidatos",
        json={
            "full_name": "Juan Perez",
            "current_position": "Gerente Comercial",
            "current_company": "Empresa Demo",
            "country": "Chile",
        },
    )
    assert candidate_response.status_code == 201
    candidate_id = candidate_response.json()["id"]

    mandate_id = _create_mandate(client)
    generated_spec = client.post(f"/api/mandatos/{mandate_id}/generar-perfil-objetivo")
    assert generated_spec.status_code == 201
    position_spec_id = generated_spec.json()["id"]

    upload_response = client.post(
        f"/api/candidatos/{candidate_id}/documentos",
        files={
            "file": (
                "juan-perez.docx",
                _create_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert upload_response.status_code == 201
    upload_body = upload_response.json()
    assert upload_body["text_extraction_status"] in {"Texto extraido", "Requiere OCR"}
    document_id = upload_body["id"]

    profile_response = client.post(f"/api/documentos-candidato/{document_id}/generar-perfil")
    assert profile_response.status_code == 201
    profile_body = profile_response.json()
    assert profile_body["candidate_id"] == candidate_id

    evaluation_response = client.post(
        "/api/evaluaciones",
        json={"candidate_id": candidate_id, "position_spec_id": position_spec_id},
    )
    assert evaluation_response.status_code == 201
    evaluation_body = evaluation_response.json()
    assert 0 <= evaluation_body["total_score"] <= 100

    list_by_candidate = client.get(f"/api/candidatos/{candidate_id}/evaluaciones")
    assert list_by_candidate.status_code == 200
    assert len(list_by_candidate.json()) == 1

    evaluation_id = evaluation_body["id"]
    word_report = client.post(f"/api/evaluaciones/{evaluation_id}/reportes/word")
    assert word_report.status_code == 200
    assert (
        word_report.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    pdf_report = client.post(f"/api/evaluaciones/{evaluation_id}/reportes/pdf")
    assert pdf_report.status_code == 200
    assert pdf_report.headers["content-type"] == "application/pdf"
